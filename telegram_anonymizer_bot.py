
"""
Telegram-бот для обезличивания юридических документов
Поддерживает форматы: PDF, DOCX
Обезличивает: ФИО, организации, ИНН, ОГРН, телефоны, email, адреса, банковские реквизиты
"""

import os
import re
import logging
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
import PyPDF2
from docx import Document
from pdf2docx import Converter
from natasha import (
    Segmenter, MorphVocab, NewsEmbedding, NewsMorphTagger, NewsNERTagger,
    Doc
)

# Настройка логирования
logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)

# Инициализация Natasha для распознавания именованных сущностей
segmenter = Segmenter()
morph_vocab = MorphVocab()
emb = NewsEmbedding()
morph_tagger = NewsMorphTagger(emb)
ner_tagger = NewsNERTagger(emb)

# Токен бота (замените на свой)
BOT_TOKEN = 'YOUR_BOT_TOKEN_HERE'

class DocumentAnonymizer:
    """Класс для обезличивания текста"""

    @staticmethod
    def anonymize_with_regex(text: str) -> str:
        """Обезличивание с помощью регулярных выражений"""

        # ИНН (10 или 12 цифр)
        text = re.sub(r'\b\d{10}(?!\d)|\b\d{12}(?!\d)', '[ИНН]', text)

        # ОГРН (13 цифр) и ОГРНИП (15 цифр)
        text = re.sub(r'\b\d{13}(?!\d)', '[ОГРН]', text)
        text = re.sub(r'\b\d{15}(?!\d)', '[ОГРНИП]', text)

        # КПП (9 цифр)
        text = re.sub(r'\b\d{9}(?!\d)', '[КПП]', text)

        # БИК (9 цифр, начинается с 04)
        text = re.sub(r'\b04\d{7}(?!\d)', '[БИК]', text)

        # Расчетный счет (20 цифр)
        text = re.sub(r'\b\d{20}(?!\d)', '[Р/С]', text)

        # Корреспондентский счет
        text = re.sub(r'\b301\d{17}(?!\d)', '[К/С]', text)

        # Телефоны
        text = re.sub(r'(\+7|8|7)?[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}', '[ТЕЛЕФОН]', text)

        # Email
        text = re.sub(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', '[EMAIL]', text)

        # Серия и номер паспорта (12 34 567890)
        text = re.sub(r'\b\d{2}\s*\d{2}\s*\d{6}\b', '[ПАСПОРТ]', text)

        # СНИЛС (123-456-789 01)
        text = re.sub(r'\b\d{3}[\-\s]?\d{3}[\-\s]?\d{3}[\s]?\d{2}\b', '[СНИЛС]', text)

        return text

    @staticmethod
    def anonymize_with_ner(text: str) -> str:
        """Обезличивание с помощью NER (Named Entity Recognition)"""
        try:
            doc = Doc(text)
            doc.segment(segmenter)
            doc.tag_morph(morph_tagger)
            doc.tag_ner(ner_tagger)

            # Сортируем сущности по убыванию позиции, чтобы замены не влияли на индексы
            entities = sorted(doc.spans, key=lambda x: x.start, reverse=True)

            text_list = list(text)
            for span in entities:
                if span.type == 'PER':  # Персона (ФИО)
                    replacement = '[ФИО]'
                elif span.type == 'ORG':  # Организация
                    replacement = '[ОРГАНИЗАЦИЯ]'
                elif span.type == 'LOC':  # Локация (адрес)
                    replacement = '[АДРЕС]'
                else:
                    continue

                # Заменяем текст
                text_list[span.start:span.stop] = replacement

            return ''.join(text_list)
        except Exception as e:
            logger.error(f"Ошибка NER: {e}")
            return text

    @staticmethod
    def full_anonymize(text: str) -> str:
        """Полное обезличивание: regex + NER"""
        # Сначала применяем regex для структурированных данных
        text = DocumentAnonymizer.anonymize_with_regex(text)
        # Затем NER для имен и организаций
        text = DocumentAnonymizer.anonymize_with_ner(text)
        return text


class FileProcessor:
    """Класс для обработки файлов"""

    @staticmethod
    def process_docx(input_path: str, output_path: str):
        """Обработка DOCX файла"""
        doc = Document(input_path)

        # Обрабатываем параграфы
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                anonymized = DocumentAnonymizer.full_anonymize(paragraph.text)
                paragraph.text = anonymized

        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        anonymized = DocumentAnonymizer.full_anonymize(cell.text)
                        cell.text = anonymized

        doc.save(output_path)

    @staticmethod
    def process_pdf(input_path: str, output_path: str):
        """Обработка PDF файла (конвертация в DOCX с обезличиванием)"""
        temp_docx = input_path.replace('.pdf', '_temp.docx')

        try:
            # Конвертируем PDF в DOCX
            cv = Converter(input_path)
            cv.convert(temp_docx, start=0, end=None)
            cv.close()

            # Обезличиваем DOCX
            FileProcessor.process_docx(temp_docx, output_path)

            # Удаляем временный файл
            if os.path.exists(temp_docx):
                os.remove(temp_docx)

        except Exception as e:
            logger.error(f"Ошибка обработки PDF: {e}")
            # Если не удалось конвертировать, используем простое извлечение текста
            FileProcessor.process_pdf_simple(input_path, output_path)

    @staticmethod
    def process_pdf_simple(input_path: str, output_path: str):
        """Простая обработка PDF (извлечение текста без сохранения форматирования)"""
        pdf_reader = PyPDF2.PdfReader(input_path)
        doc = Document()

        for page in pdf_reader.pages:
            text = page.extract_text()
            if text.strip():
                anonymized = DocumentAnonymizer.full_anonymize(text)
                doc.add_paragraph(anonymized)

        doc.save(output_path)


# Обработчики команд бота
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик команды /start"""
    welcome_text = """
🔒 *Бот для обезличивания юридических документов*

Я помогу вам автоматически удалить персональные данные из документов.

*Что я умею обезличивать:*
• ФИО и инициалы
• Названия организаций
• ИНН, ОГРН, ОГРНИП, КПП
• Банковские реквизиты (БИК, Р/С, К/С)
• Телефоны и email
• Паспортные данные и СНИЛС
• Адреса

*Поддерживаемые форматы:*
• DOCX (MS Word)
• PDF

*Как использовать:*
1. Отправьте мне документ (PDF или DOCX)
2. Дождитесь обработки
3. Получите обезличенный документ в формате DOCX

⚠️ *Внимание:* Бот работает локально, ваши файлы удаляются сразу после обработки.

Отправьте документ для начала работы!
    """
    await update.message.reply_text(welcome_text, parse_mode='Markdown')


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик команды /help"""
    help_text = """
📖 *Справка по использованию*

*Поддерживаемые форматы:*
• DOCX - сохраняется структура документа
• PDF - конвертируется в DOCX с обезличиванием

*Что обезличивается:*
• Персональные данные: ФИО, паспорта, СНИЛС
• Организации: названия компаний
• Реквизиты: ИНН, ОГРН, КПП, БИК, счета
• Контакты: телефоны, email
• Адреса и локации

*Ограничения:*
• Максимальный размер файла: 20 МБ
• Качество распознавания зависит от структуры документа
• PDF-файлы конвертируются в DOCX

*Безопасность:*
Все файлы автоматически удаляются после обработки. Данные не хранятся.
    """
    await update.message.reply_text(help_text, parse_mode='Markdown')


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик загруженных документов"""
    document = update.message.document

    # Проверяем формат файла
    if not (document.file_name.endswith('.pdf') or document.file_name.endswith('.docx')):
        await update.message.reply_text(
            "❌ Неподдерживаемый формат файла. Пожалуйста, отправьте PDF или DOCX документ."
        )
        return

    await update.message.reply_text("⏳ Обрабатываю документ... Это может занять некоторое время.")

    try:
        # Скачиваем файл
        file = await context.bot.get_file(document.file_id)
        input_path = f"input_{document.file_name}"
        await file.download_to_drive(input_path)

        # Определяем выходной файл
        output_filename = document.file_name.replace('.pdf', '_anonymized.docx').replace('.docx', '_anonymized.docx')
        output_path = f"output_{output_filename}"

        # Обрабатываем файл
        if document.file_name.endswith('.pdf'):
            FileProcessor.process_pdf(input_path, output_path)
        else:
            FileProcessor.process_docx(input_path, output_path)

        # Отправляем результат
        with open(output_path, 'rb') as f:
            await update.message.reply_document(
                document=f,
                filename=output_filename,
                caption="✅ Документ успешно обезличен!\n\n⚠️ Проверьте результат перед использованием."
            )

        # Удаляем временные файлы
        if os.path.exists(input_path):
            os.remove(input_path)
        if os.path.exists(output_path):
            os.remove(output_path)

    except Exception as e:
        logger.error(f"Ошибка обработки документа: {e}")
        await update.message.reply_text(
            f"❌ Произошла ошибка при обработке документа: {str(e)}\n\n"
            "Пожалуйста, попробуйте еще раз или отправьте другой файл."
        )

        # Очистка файлов в случае ошибки
        try:
            if os.path.exists(input_path):
                os.remove(input_path)
            if os.path.exists(output_path):
                os.remove(output_path)
        except:
            pass


async def error_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик ошибок"""
    logger.error(f"Update {update} caused error {context.error}")


def main():
    """Запуск бота"""
    # Создаем приложение
    application = Application.builder().token(BOT_TOKEN).build()

    # Регистрируем обработчики
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    application.add_error_handler(error_handler)

    # Запускаем бота
    logger.info("Бот запущен!")
    application.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == '__main__':
    main()
